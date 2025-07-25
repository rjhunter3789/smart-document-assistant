import os
import io
import fitz  # PyMuPDF
from docx import Document as DocxDoc
import streamlit as st
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.auth.transport.requests import Request
from datetime import datetime

# Config
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
DOCS_DIR = "docs"
os.makedirs(DOCS_DIR, exist_ok=True)

# Mobile-optimized page configuration
st.set_page_config(
    page_title="🎤 Voice-Powered Document Assistant",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="collapsed",
    menu_items={
        'Get Help': 'https://your-help-url.com',
        'Report a bug': "https://your-bug-report-url.com",
        'About': "# Voice-Powered Smart Document Assistant\nMobile-optimized AI-powered document search with voice commands"
    }
)

# Custom CSS for mobile optimization + Voice UI
st.markdown("""
<style>
    /* Mobile-first responsive design */
    .main .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
        padding-left: 1rem;
        padding-right: 1rem;
        max-width: 100%;
    }
    
    /* Header styling */
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 15px;
        margin-bottom: 1.5rem;
        text-align: center;
        color: white;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .main-header h1 {
        margin: 0;
        font-size: 1.8rem;
        font-weight: 700;
    }
    
    .main-header p {
        margin: 0.5rem 0 0 0;
        opacity: 0.9;
        font-size: 1rem;
    }
    
    /* Voice Interface Styling */
    .voice-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 20px;
        padding: 2rem;
        margin: 1.5rem 0;
        text-align: center;
        color: white;
        box-shadow: 0 8px 25px rgba(102, 126, 234, 0.3);
    }
    
    /* File cards styling */
    .file-card {
        background: white;
        border: 1px solid #e1e5e9;
        border-radius: 12px;
        padding: 1rem;
        margin: 0.75rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        transition: all 0.3s ease;
    }
    
    .file-card:hover {
        box-shadow: 0 4px 16px rgba(0,0,0,0.12);
        transform: translateY(-2px);
    }
    
    .file-name {
        font-weight: 600;
        color: #2c3e50;
        font-size: 1rem;
        margin-bottom: 0.5rem;
        word-break: break-word;
    }
    
    /* Navigation section */
    .nav-section {
        background: #f8f9fa;
        border-radius: 12px;
        padding: 1rem;
        margin: 1rem 0;
        border-left: 4px solid #667eea;
    }
    
    .nav-title {
        font-weight: 600;
        color: #2c3e50;
        margin-bottom: 0.5rem;
        font-size: 1rem;
    }
    
    /* Results section */
    .results-header {
        background: linear-gradient(90deg, #56ccf2, #2f80ed);
        color: white;
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
        text-align: center;
        font-weight: 600;
    }
    
    /* Status messages */
    .status-success {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        padding: 0.75rem;
        border-radius: 8px;
        margin: 0.5rem 0;
    }
    
    .status-warning {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        color: #856404;
        padding: 0.75rem;
        border-radius: 8px;
        margin: 0.5rem 0;
    }
    
    .status-info {
        background: #d1ecf1;
        border: 1px solid #bee5eb;
        color: #0c5460;
        padding: 0.75rem;
        border-radius: 8px;
        margin: 0.5rem 0;
    }
    
    /* Mobile-friendly dropdowns */
    .stSelectbox > div > div {
        font-size: 16px !important;
        min-height: 44px;
    }
    
    /* Search input styling */
    .stTextInput > div > div > input {
        border-radius: 25px;
        border: 2px solid #e1e5e9;
        padding: 0.75rem 1rem;
        font-size: 16px;
        min-height: 44px;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 0.2rem rgba(102,126,234,0.25);
    }
    
    /* Responsive font sizes */
    @media (max-width: 768px) {
        .main-header h1 {
            font-size: 1.5rem;
        }
        
        .main-header p {
            font-size: 0.9rem;
        }
        
        .file-name {
            font-size: 0.95rem;
        }
    }
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# Authentication function
def authenticate_gdrive():
    """Authenticate with Google Drive using Streamlit secrets"""
    try:
        from google.oauth2.credentials import Credentials
        
        # Create credentials from Streamlit secrets
        creds_info = {
            "client_id": st.secrets["google"]["client_id"],
            "client_secret": st.secrets["google"]["client_secret"],
            "refresh_token": st.secrets["google"]["refresh_token"],
            "token": st.secrets["google"]["token"],
            "token_uri": st.secrets["google"]["token_uri"],
        }
        
        # Create credentials object
        creds = Credentials.from_authorized_user_info(creds_info, SCOPES)
        
        # Refresh if needed
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
        
        # Build and return the service
        return build('drive', 'v3', credentials=creds)
        
    except Exception as e:
        st.error(f"🚫 Authentication failed: {str(e)}")
        return None

def get_all_folders_recursive(service, parent_id, parent_path=""):
    """Get ALL folders - no limits, full depth for business use"""
    folders = []
    
    try:
        query = f"'{parent_id}' in parents and mimeType = 'application/vnd.google-apps.folder' and trashed = false"
        results = service.files().list(q=query, fields="files(id, name)", pageSize=1000).execute()
        items = results.get("files", [])

        for item in items:
            full_path = f"{parent_path}/{item['name']}".strip("/")
            folders.append({
                "id": item["id"],
                "name": item["name"],
                "full_path": full_path
            })
            
            # FULL RECURSION - No limits for business requirements
            subfolders = get_all_folders_recursive(service, item["id"], full_path)
            folders.extend(subfolders)
                
    except Exception as e:
        st.error(f"Error accessing folder {parent_path}: {str(e)}")
    
    return folders

def list_files(service, folder_id):
    """List files in a specific folder"""
    try:
        results = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields="files(id, name, mimeType, modifiedTime, size)"
        ).execute()
        files = results.get('files', [])
        return files
    except Exception as e:
        st.error(f"Error listing files: {str(e)}")
        return []

def download_file(service, file_id, name):
    """Download a file from Google Drive"""
    try:
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        path = os.path.join(DOCS_DIR, name)
        with open(path, 'wb') as f:
            f.write(fh.getvalue())
        return path
    except Exception as e:
        st.error(f"Error downloading file: {str(e)}")
        return None

def extract_text_from_file(service, file_id, file_name):
    """Extract text content from a Google Drive file"""
    try:
        # Download the file
        request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
               _, done = downloader.next_chunk()
        
        # Extract text based on file type
        file_content = fh.getvalue()
        
        if file_name.lower().endswith('.pdf'):
            # Extract text from PDF
            pdf_doc = fitz.open(stream=file_content, filetype="pdf")
            text = ""
            for page in pdf_doc:
                text += page.get_text() + "\n"
            pdf_doc.close()
            return text
            
        elif file_name.lower().endswith('.docx'):
            # Extract text from Word document
            doc = DocxDoc(io.BytesIO(file_content))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
            
        elif file_name.lower().endswith('.txt'):
            # Extract text from text file
            return file_content.decode('utf-8')
            
        elif file_name.lower().endswith(('.xlsx', '.xls')):
            # Extract readable summary from Excel file
            import pandas as pd
            
            try:
                # Read Excel file
                excel_file = pd.ExcelFile(io.BytesIO(file_content))
                summary_parts = []
                
                summary_parts.append(f"Excel File: {file_name}")
                summary_parts.append(f"Number of sheets: {len(excel_file.sheet_names)}")
                summary_parts.append("Sheet contents:")
                
                # Process each sheet (limit to first 3 sheets for mobile)
                for i, sheet_name in enumerate(excel_file.sheet_names[:3]):
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    
                    summary_parts.append(f"\nSheet '{sheet_name}':")
                    summary_parts.append(f"  - {df.shape[0]} rows, {df.shape[1]} columns")
                    
                    # Add column names
                    if len(df.columns) > 0:
                        columns = list(df.columns)[:5]  # First 5 columns
                        summary_parts.append(f"  - Columns: {', '.join(str(col) for col in columns)}")
                        if len(df.columns) > 5:
                            summary_parts.append(f"    ... and {len(df.columns) - 5} more columns")
                    
                    # Add sample data from first few rows
                    if not df.empty:
                        summary_parts.append("  - Sample data:")
                        for idx, row in df.head(3).iterrows():
                            row_summary = []
                            for col in df.columns[:3]:  # First 3 columns only
                                value = str(row[col])[:30]  # Truncate long values
                                if len(str(row[col])) > 30:
                                    value += "..."
                                row_summary.append(f"{col}: {value}")
                            summary_parts.append(f"    Row {idx + 1}: {', '.join(row_summary)}")
                
                if len(excel_file.sheet_names) > 3:
                    summary_parts.append(f"\n... and {len(excel_file.sheet_names) - 3} more sheets")
                
                return "\n".join(summary_parts)
                
            except Exception as excel_error:
                return f"Excel file detected but could not be processed: {str(excel_error)}"
        
        else:
            return "Text extraction not supported for this file type."
            
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def smart_text_truncate(text, max_length=500):
    """Intelligently truncate text at sentence boundaries for natural speech"""
    if not text or len(text) <= max_length:
        return text
    
    # Clean up the text first
    import re
    text = re.sub(r'\s+', ' ', text).strip()  # Remove extra whitespace
    text = re.sub(r'[^\w\s.,!?;:-]', '', text)  # Remove special chars but keep punctuation
    
    # If text is short enough, return as-is
    if len(text) <= max_length:
        return text
    
    # Find the best place to cut - look for sentence endings near the limit
    truncated = text[:max_length]
    
    # Look for sentence endings (. ! ?) in the last 100 characters
    sentence_endings = []
    for i, char in enumerate(truncated):
        if char in '.!?':
            sentence_endings.append(i)
    
    if sentence_endings:
        # Cut at the last complete sentence
        cut_point = sentence_endings[-1] + 1
        result = text[:cut_point].strip()
        if len(result) > 50:  # Make sure we don't get too short
            return result
    
    # If no good sentence ending, look for clause endings (, ; :)
    clause_endings = []
    for i, char in enumerate(truncated):
        if char in ',;:':
            clause_endings.append(i)
    
    if clause_endings:
        # Cut at the last clause ending
        cut_point = clause_endings[-1] + 1
        result = text[:cut_point].strip()
        if len(result) > 100:  # Make sure we have substantial content
            return result
    
    # Last resort: cut at word boundary
    words = truncated.split(' ')
    if len(words) > 1:
        # Remove the last incomplete word
        complete_words = ' '.join(words[:-1])
        if len(complete_words) > 50:
            return complete_words + "..."
    
    # Fallback: just truncate and add ellipsis
    return truncated + "..."

def generate_summary(text):
    """Generate a simple summary (lightweight version)"""
    if not text.strip():
        return "No readable content found."
    
    # Simple text truncation for demo
    sentences = text.split('.')[:3]  # First 3 sentences
    summary = '. '.join(sentences).strip()
    
    if len(summary) > 300:
        summary = summary[:300] + "..."
    
    return summary if summary else "Document summary would appear here with full AI capabilities."

def process_voice_command(docs, service):
    """Process voice commands and execute actions"""
    if 'voice_command' in st.session_state and st.session_state.voice_command:
        command = st.session_state.voice_command
        action = command['action']
        search_term = command['search_term']
        
        # Search for matching documents
        matching_docs = []
        for doc in docs:
            doc_name = doc.get('name', '').lower()
            if search_term in doc_name:
                matching_docs.append(doc)
        
        if matching_docs:
            st.success(f"🎯 Found {len(matching_docs)} document(s) matching '{search_term}'")
            
            # Process the first matching document
            target_doc = matching_docs[0]
            doc_name = target_doc.get('name', 'Unknown Document')
            
            if action == 'summarize':
                st.markdown("### 🤖 Auto-Generated Summary")
                
                # Real document summary
                with st.spinner(f"Analyzing {doc_name}..."):
                    file_text = extract_text_from_file(service, target_doc['id'], doc_name)
                    if file_text and not file_text.startswith("Error"):
                        summary = generate_summary(file_text)
                        st.info(f"📋 **Summary of {doc_name}:**\n\n{summary}")
                        
                        # Auto-read the summary using speech synthesis
                        st.components.v1.html(f"""
                        <script>
                        setTimeout(function() {{
                            if ('speechSynthesis' in window) {{
                                const utterance = new SpeechSynthesisUtterance('{summary.replace("'", "\\'")}');
                                utterance.rate = 0.8;
                                utterance.pitch = 1.0;
                                utterance.volume = 1.0;
                                speechSynthesis.speak(utterance);
                            }}
                        }}, 1000);
                        </script>
                        """, height=0)
                    else:
                        st.error("Unable to analyze document")
            
            elif action == 'read':
                st.markdown("### 🔊 Auto-Reading Document")
                
                with st.spinner(f"Extracting text from {doc_name}..."):
                    file_text = extract_text_from_file(service, target_doc['id'], doc_name)
                    if file_text and not file_text.startswith("Error"):
                        # Truncate for speech
                        speech_text = smart_text_truncate(file_text, 600)
                        st.info(f"🔊 Reading: {doc_name}")
                        
                        # Auto-read the document
                        st.components.v1.html(f"""
                        <script>
                        setTimeout(function() {{
                            if ('speechSynthesis' in window) {{
                                const utterance = new SpeechSynthesisUtterance('{speech_text.replace("'", "\\'")}');
                                utterance.rate = 0.8;
                                utterance.pitch = 1.0;
                                utterance.volume = 1.0;
                                speechSynthesis.speak(utterance);
                            }}
                        }}, 1000);
                        </script>
                        """, height=0)
                    else:
                        st.error("Unable to extract text for reading")
            
            elif action == 'find':
                st.markdown("### 🔍 Search Results")
                st.success(f"Found {len(matching_docs)} document(s):")
                for doc in matching_docs:
                    st.write(f"📄 {doc.get('name', 'Unknown')}")
            
            # Clear the command after processing
            st.session_state.voice_command = None
            
        else:
            st.warning(f"❌ No documents found matching '{search_term}'. Try a different search term.")

# Main UI
st.markdown("""
<div class="main-header">
    <h1>🎤 Voice-Powered Document Assistant</h1>
    <p>AI-powered mobile document search with voice commands</p>
</div>
""", unsafe_allow_html=True)

# Voice Interface Section with Audio Permission Request
st.markdown("""
<div class="voice-container">
    <h3 style="margin-top: 0;">🎙️ Voice Commands</h3>
    <p>Use the voice input below to search your documents</p>
</div>
""", unsafe_allow_html=True)

# Audio permission and test section
st.markdown("### 🔊 Audio Setup")
col1, col2 = st.columns([1, 1])

with col1:
    if st.button("🔧 Test Audio System", use_container_width=True):
        st.markdown("""
        <script>
        // Request audio permission and test
        async function testAudioPermission() {
            try {
                // Test speech synthesis
                if ('speechSynthesis' in window) {
                    const testUtterance = new SpeechSynthesisUtterance('Audio test successful. Text to speech is working.');
                    testUtterance.volume = 1;
                    testUtterance.rate = 0.8;
                    speechSynthesis.speak(testUtterance);
                    
                    // Show success message
                    const statusDiv = document.createElement('div');
                    statusDiv.style.cssText = 'background: #d4edda; padding: 10px; border-radius: 5px; margin: 10px 0; color: #155724;';
                    statusDiv.innerHTML = '✅ Audio test started! You should hear "Audio test successful..."';
                    document.body.appendChild(statusDiv);
                    
                    setTimeout(() => {
                        if (statusDiv.parentNode) {
                            statusDiv.parentNode.removeChild(statusDiv);
                        }
                    }, 5000);
                    
                } else {
                    alert('Speech synthesis not supported in this browser. Try Chrome or Edge.');
                }
            } catch (error) {
                alert('Audio test failed: ' + error.message);
            }
        }
        
        testAudioPermission();
        </script>
        """, unsafe_allow_html=True)

with col2:
    if st.button("📋 Audio Troubleshooting", use_container_width=True):
        st.markdown("""
        **🔧 If you can't hear audio:**
        
        1. **Check browser settings:**
           - Chrome: Settings → Site Settings → Sound → Allow
           - Edge: Settings → Site permissions → Sound → Allow
        
        2. **Try these steps:**
           - Refresh the page
           - Try an incognito/private window
           - Check your system volume
           - Try a different browser (Chrome works best)
        
        3. **Manual test:**
           - Open browser console (F12)
           - Type: `speechSynthesis.speak(new SpeechSynthesisUtterance("test"))`
           - Press Enter
        """)

# Add clear folder cache button for troubleshooting
if st.button("🔄 Clear Cache & Reload Folders"):
    st.session_state.all_folders = None
    st.session_state.folder_cache_time = None
    st.rerun()

# Voice input with enhanced voice command processing
st.markdown("### 🎤 Voice Commands")

# Voice command help
with st.expander("🎯 Voice Command Examples"):
    st.markdown("""
    **🚗 Hands-Free Commands:**
    - *"Summarize the Chat AI file"*
    - *"Read me the Impel document"* 
    - *"Find and summarize automotive platform"*
    - *"Review the FAQ handout"*
    
    **📱 How to Use:**
    1. **Tap the voice box** below
    2. **Use your keyboard's mic** button 🎤
    3. **Speak your command** clearly
    4. **App will auto-execute** the command
    """)

voice_input = st.text_input(
    "🎤 Voice Command", 
    placeholder="Say: 'Summarize the Chat AI file' or 'Read me the automotive platform document'",
    help="💡 Speak commands like 'summarize [filename]' or 'read [filename]'",
    label_visibility="collapsed"
)

# Process voice commands
if voice_input:
    st.success(f"🎤 Voice command: **{voice_input}**")
    
    # Parse voice command
    command_lower = voice_input.lower()
    
    # Extract action and filename
    action = None
    search_term = None
    
    if any(word in command_lower for word in ['summarize', 'summary', 'review', 'tell me about', 'what is']):
        action = 'summarize'
    elif any(word in command_lower for word in ['read', 'read back', 'read to me', 'play', 'speak']):
        action = 'read'
    elif any(word in command_lower for word in ['find', 'search', 'show', 'look for', 'get me']):
        action = 'find'
    
    # Extract search terms (remove command words)
    search_words = command_lower
    for remove_word in ['summarize', 'summary', 'review', 'read', 'back', 'to', 'me', 'the', 'file', 'on', 'about', 'find', 'search', 'show', 'and']:
        search_words = search_words.replace(remove_word, ' ')
    
    search_term = ' '.join(search_words.split()).strip()
    
    if action and search_term:
        st.info(f"🎯 **Command understood:** {action.title()} documents matching '{search_term}'")
        
        # Store voice command in session state for processing later
        if 'voice_command' not in st.session_state:
            st.session_state.voice_command = {}
        
        st.session_state.voice_command = {
            'action': action,
            'search_term': search_term,
            'original_command': voice_input
        }
    else:
        st.warning("🤔 Command not understood. Try: 'Summarize [filename]' or 'Read [filename]'")

# Search interface
st.markdown("### 🔍 Search Documents")
filename_query = st.text_input(
    "📄 Search by filename", 
    placeholder="Type filename to search...", 
    label_visibility="collapsed"
)

# Use the combined search query
active_query = voice_input if voice_input else filename_query

# Initialize session state for caching
if 'all_folders' not in st.session_state:
    st.session_state.all_folders = None
if 'folder_cache_time' not in st.session_state:
    st.session_state.folder_cache_time = None
if 'current_step' not in st.session_state:
    st.session_state.current_step = 1
if 'selected_folder_id' not in st.session_state:
    st.session_state.selected_folder_id = None
if 'selected_folder_path' not in st.session_state:
    st.session_state.selected_folder_path = None

# Authentication and Google Drive setup
try:
    service = authenticate_gdrive()
    
    if service:
        st.success("🔗 Connected to Google Drive! Loading your folders...")
    else:
        st.error("🚫 Google Drive not connected. Please check your authentication credentials.")
        
except Exception as e:
    st.error(f"🚫 Authentication setup failed: {str(e)}")
    service = None

# Google Drive folder setup
root_id = "1galnuNa9g7xoULx3Ka8vs79-NuJUA4n6"  # Your root folder ID
all_folders = []
docs = []

if service:
    # Get all folders (with smart caching)
    try:
        docs = []  # Initialize as empty list

        # Check if we need to refresh cache (20 minute expiry - balance between freshness and load time)
        import time
        current_time = time.time()
        cache_expired = (
            st.session_state.folder_cache_time is None or 
            current_time - st.session_state.folder_cache_time > 1200  # 20 minutes - longer cache for full load
        )
        
        if st.session_state.all_folders is None or cache_expired:
            with st.spinner("📁 Loading complete folder structure (this may take 30-60 seconds for 38 dealers)..."):
                st.session_state.all_folders = get_all_folders_recursive(service, root_id)
                st.session_state.folder_cache_time = current_time
                st.success(f"📁 Complete! Loaded {len(st.session_state.all_folders)} folders from all dealers")
        else:
            st.info(f"📁 Using cached structure ({len(st.session_state.all_folders)} folders) - Refreshes every 20 minutes")
            
        all_folders = st.session_state.all_folders
        
        if all_folders:
            st.success(f"📁 Ready! {len(all_folders)} folders available")
            
            # Filter out WMA Test and excluded paths
            excluded_names = ["WMA Test"]
            excluded_paths = ["WMA RAG/WMA Team/RAG_App"]
            
            filtered_folders = [
                f for f in all_folders 
                if f['name'] not in excluded_names and 
                not any(f['full_path'].startswith(p) for p in excluded_paths)
            ]
            
            # Find team folders and WMA folders for step 1
            team_names = ["Aaron", "Brody", "Dona", "Eric", "Grace", "Jeff", "Jessica", "Jill", "John", "Jon", "Kirk", "Owen", "Paul"]
            wma_folders = [f for f in filtered_folders if 'WMA' in f['name']]
            team_folders = [f for f in filtered_folders if f['name'] in team_names]
            
            # Combine and get unique main folder names
            main_folders = wma_folders + team_folders
            main_folder_names = list(set([f['name'] for f in main_folders if '/' not in f['full_path']]))
            main_folder_names.sort()
            
            # Folder Navigation with Security
            st.markdown("### 👤 User Identity")
            
            # Step 0: User Authentication
            user_identity = st.selectbox(
                "**Who are you?**",
                ["Aaron", "Brody", "Dona", "Eric", "Grace", "Jeff", "Jessica", "Jill", "John", "Jon", "Kirk", "Owen", "Paul"],
                label_visibility="visible"
            )
            
            st.markdown("### 📁 Browse Your Folders")
            
            # Filter folders based on user identity - security layer
            if user_identity:
                allowed_folders = [user_identity, "WMA Team"]  # User can see their own folder + shared WMA folders
                
                # Filter main folders to only show allowed ones
                accessible_main_folders = [name for name in main_folder_names if any(name.startswith(allowed) for allowed in allowed_folders)]
                
                if not accessible_main_folders:
                    st.error(f"No folders found for {user_identity}. Please contact your administrator.")
                    docs = []
                else:
                    # Step 1: Select main folder (filtered by user identity)
                    selected_main_folder = st.selectbox(
                        f"**Step 1:** Choose your folder (showing folders for {user_identity})", 
                        accessible_main_folders,
                        label_visibility="visible"
                    )
                    
                    # Step 2: Get subfolders for selected main folder (with security check)
                    if selected_main_folder in accessible_main_folders:
                        subfolders = [f for f in filtered_folders if f['full_path'].startswith(selected_main_folder + '/')]
                        level_1_subfolders = [f for f in subfolders if f['full_path'].count('/') == 1]
                        
                        selected_folder = None
                        selected_folder_path = selected_main_folder
                        
                        if level_1_subfolders:
                            level_1_paths = [f['full_path'] for f in level_1_subfolders]
                            level_1_paths.sort()
                            
                            selected_level_1 = st.selectbox(
                                "**Step 2:** Choose subfolder", 
                                level_1_paths,
                                label_visibility="visible"
                            )
                            
                            # Step 3: Get level 2 subfolders
                            level_2_subfolders = [f for f in filtered_folders if f['full_path'].startswith(selected_level_1 + '/')]
                            immediate_level_2 = [f for f in level_2_subfolders if f['full_path'].count('/') == 2]
                            
                            if immediate_level_2:
                                level_2_paths = [f['full_path'] for f in immediate_level_2]
                                level_2_paths.sort()
                                
                                selected_level_2 = st.selectbox(
                                    "**Step 3:** Choose final folder", 
                                    level_2_paths,
                                    label_visibility="visible"
                                )
                                
                                selected_folder = next((f for f in immediate_level_2 if f['full_path'] == selected_level_2), None)
                                selected_folder_path = selected_level_2
                            else:
                                # Use level 1 as final selection
                                selected_folder = next((f for f in level_1_subfolders if f['full_path'] == selected_level_1), None)
                                selected_folder_path = selected_level_1
                        else:
                            # Use main folder as final selection
                            selected_folder = next((f for f in main_folders if f['name'] == selected_main_folder and '/' not in f['full_path']), None)
                            selected_folder_path = selected_main_folder
                    else:
                        st.error(f"🚫 Access denied: {user_identity} cannot access {selected_main_folder} folder")
                        docs = []
                        selected_folder = None
                    
                    # Load files with security validation
                    if selected_folder:
                        # Double-check security: ensure user can access this folder
                        folder_owner = selected_folder['full_path'].split('/')[0]
                        if folder_owner in allowed_folders or any(selected_folder['full_path'].startswith(allowed) for allowed in allowed_folders):
                            with st.spinner(f"Loading files from {selected_folder['name']}..."):
                                docs = list_files(service, selected_folder['id'])
                            
                            if docs:
                                st.success(f"📄 Found **{len(docs)}** files in **{selected_folder['name']}**")
                            else:
                                st.warning(f"No files found in {selected_folder['full_path']}. Folder appears to be empty.")
                                docs = []
                        else:
                            st.error(f"🚫 Security violation: {user_identity} attempted to access unauthorized folder {selected_folder['full_path']}")
                            docs = []
                    else:
                        st.warning("Please select a folder to view its contents.")
                        docs = []
            else:
                st.warning("Please select your identity to continue.")
                docs = []
        else:
            st.warning("No folders found in Google Drive.")
            docs = []
            
    except Exception as e:
        st.error(f"Error loading Google Drive folders: {str(e)}")
        docs = []
else:
    st.error("Please authenticate with Google Drive to continue.")
    docs = []

# Process voice commands if any
if docs:
    process_voice_command(docs, service)

# Apply search filter
if active_query and docs:
    original_count = len(docs)
    filtered_docs = [doc for doc in docs if active_query.lower() in doc["name"].lower()]
    if filtered_docs:
        docs = filtered_docs
        st.info(f"🔍 **Search Results:** Found {len(docs)} documents matching '{active_query}' (out of {original_count} total)")
    else:
        st.warning(f"❌ No documents found matching '{active_query}' in the selected folder")

# Display results
st.markdown("---")
st.markdown("### 📄 Documents")

if docs:
    st.success(f"**{len(docs)} documents** from **{selected_folder_path}**")
    
    # Display each document
    for i, doc in enumerate(docs):
        doc_name = doc.get('name', 'Unknown Document')
        
        st.markdown(f"""
        <div class="file-card">
            <div class="file-name">📄 {doc_name}</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Action buttons
        col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
        
        with col1:
            if st.button("📖 Summary", key=f"sum_{i}", use_container_width=True):
                # Real document summary
                if 'id' in doc:
                    with st.spinner(f"Analyzing {doc_name}..."):
                        file_text = extract_text_from_file(service, doc['id'], doc_name)
                        if file_text and not file_text.startswith("Error"):
                            summary = generate_summary(file_text)
                            st.success(f"**📋 AI Summary:** {summary}")
                            
                            # Auto-read the summary
                            st.components.v1.html(f"""
                            <script>
                            setTimeout(function() {{
                                if ('speechSynthesis' in window) {{
                                    const utterance = new SpeechSynthesisUtterance('{summary.replace("'", "\\'")}');
                                    utterance.rate = 0.8;
                                    utterance.pitch = 1.0;
                                    utterance.volume = 1.0;
                                    speechSynthesis.speak(utterance);
                                }}
                            }}, 1000);
                            </script>
                            """, height=0)
                        else:
                            st.error("Unable to analyze document")
                else:
                    st.error("Unable to analyze: File ID not available")
        
        with col2:
            if st.button("🔊 Read", key=f"speak_{i}", use_container_width=True):
                # Real document text-to-speech
                if 'id' in doc:
                    with st.spinner("Extracting text for speech..."):
                        file_text = extract_text_from_file(service, doc['id'], doc_name)
                        if file_text and not file_text.startswith("Error"):
                            # Use smart truncation to end at complete sentences
                            speech_text = smart_text_truncate(file_text, 600)
                            clean_text = speech_text.replace('"', '\\"').replace("'", "\\'").replace('\n', ' ')
                            
                            st.markdown(f"""
                            <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 1.5rem; border-radius: 12px; margin: 1rem 0; color: white;">
                                <h4 style="margin-top: 0; color: white;">🔊 Now Reading: {doc_name}</h4>
                                <p style="background: rgba(255,255,255,0.1); padding: 1rem; border-radius: 6px; font-style: italic; margin: 1rem 0;">
                                    "{speech_text}"
                                </p>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Add stop/pause controls
                            col_stop, col_pause = st.columns([1, 1])
                            with col_stop:
                                if st.button("⏹️ Stop Speech", key=f"stop_real_{i}", use_container_width=True):
                                    st.components.v1.html("""
                                    <script>
                                    if ('speechSynthesis' in window) {
                                        speechSynthesis.cancel();
                                    }
                                    document.body.innerHTML = '<div style="background: #dc3545; color: white; padding: 10px; text-align: center; border-radius: 5px;">⏹️ Document Speech Stopped</div>';
                                    </script>
                                    """, height=50)
                            
                            with col_pause:
                                if st.button("⏸️ Pause/Resume", key=f"pause_real_{i}", use_container_width=True):
                                    st.components.v1.html("""
                                    <script>
                                    if ('speechSynthesis' in window) {
                                        if (speechSynthesis.speaking && !speechSynthesis.paused) {
                                            speechSynthesis.pause();
                                            document.body.innerHTML = '<div style="background: #ffc107; color: black; padding: 10px; text-align: center; border-radius: 5px;">⏸️ Document Paused</div>';
                                        } else if (speechSynthesis.paused) {
                                            speechSynthesis.resume();
                                            document.body.innerHTML = '<div style="background: #28a745; color: white; padding: 10px; text-align: center; border-radius: 5px;">▶️ Document Resumed</div>';
                                        }
                                    }
                                    </script>
                                    """, height=50)
                            
                            st.components.v1.html(f"""
                            <script>
                            window.addEventListener('load', function() {{
                                setTimeout(function() {{
                                    if ('speechSynthesis' in window) {{
                                        let voices = speechSynthesis.getVoices();
                                        
                                        if (voices.length === 0) {{
                                            speechSynthesis.onvoiceschanged = function() {{
                                                voices = speechSynthesis.getVoices();
                                                speakDocument(voices);
                                            }};
                                        }} else {{
                                            speakDocument(voices);
                                        }}
                                        
                                        function speakDocument(voices) {{
                                            // Premium voice selection for documents
                                            const preferredVoices = [
                                                'Microsoft Zira - English (United States)',
                                                'Microsoft David - English (United States)', 
                                                'Alex', 'Samantha', 'Victoria',
                                                'Google US English', 'Ava (Enhanced)',
                                                'en-US-Wavenet'
                                            ];
                                            
                                            let selectedVoice = null;
                                            for (let preferred of preferredVoices) {{
                                                selectedVoice = voices.find(voice => 
                                                    voice.name.includes(preferred) || voice.name === preferred
                                                );
                                                if (selectedVoice) break;
                                            }}
                                            
                                            // Fallback to best available English voice
                                            if (!selectedVoice) {{
                                                selectedVoice = voices.find(voice => 
                                                    voice.lang.includes('en-US') && !voice.name.toLowerCase().includes('compact')
                                                ) || voices.find(voice => voice.lang.includes('en')) || voices[0];
                                            }}
                                            
                                            const utterance = new SpeechSynthesisUtterance(`{clean_text}`);
                                            
                                            if (selectedVoice) {{
                                                utterance.voice = selectedVoice;
                                            }}
                                            
                                            // Professional speech settings for documents
                                            utterance.rate = 0.8;       // Slower for business documents
                                            utterance.pitch = 1.0;     // Natural pitch
                                            utterance.volume = 1.0;    // Full volume
                                            
                                            speechSynthesis.speak(utterance);
                                            
                                            document.body.style.background = 'linear-gradient(45deg, #28a745, #20c997)';
                                            document.body.style.color = 'white';
                                            document.body.style.padding = '20px';
                                            document.body.style.borderRadius = '10px';
                                            document.body.innerHTML = `
                                                <h3>🔊 Reading Document: {doc_name}</h3>
                                                <p>High-Quality Voice: ${{selectedVoice ? selectedVoice.name : 'Default'}}</p>
                                                <p>Playing professional document audio...</p>
                                            `;
                                        }}
                                    }}
                                }}, 500);
                            }});
                            </script>
                            <div style="background: #28a745; color: white; padding: 10px; text-align: center; border-radius: 5px;">
                                🔊 Loading premium voice for document... Enhanced audio quality
                            </div>
                            """, height=100)
                        else:
                            st.error("Unable to extract text for speech synthesis")
                else:
                    st.error("Unable to read file: File ID not available")
        
        with col3:
            if st.button("⬇️ Download", key=f"dl_{i}", use_container_width=True):
                # Real download functionality
                if 'id' in doc:
                    with st.spinner("Downloading..."):
                        file_path = download_file(service, doc['id'], doc_name)
                        if file_path:
                            st.success(f"✅ **Downloaded:** {doc_name}")
                            with open(file_path, 'rb') as f:
                                st.download_button(
                                    label=f"💾 Save {doc_name}",
                                    data=f.read(),
                                    file_name=doc_name,
                                    key=f"download_btn_{i}"
                                )
                else:
                    st.error("Unable to download: File ID not available")
        
        with col4:
            if st.button("🔗 Share", key=f"share_{i}", use_container_width=True):
                if 'id' in doc:
                    share_url = f"https://drive.google.com/file/d/{doc['id']}/view"
                    st.success(f"**🔗 Google Drive Link:** {share_url}")
                    st.code(share_url, language=None)
                else:
                    st.error("Unable to generate share link")
        
        st.markdown("---")
else:
    st.warning("📄 No documents found in the selected location.")
    st.info("💡 Try selecting a different folder or check if the folder contains any files.")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; font-size: 0.9rem; padding: 1rem;'>
    🎤 Voice-Powered Smart Document Assistant<br>
    Powered by AI • Built for productivity • Voice-enabled<br>
    <small>💡 Tip: Use Chrome or Edge for best voice experience</small>
</div>
""", unsafe_allow_html=True)
